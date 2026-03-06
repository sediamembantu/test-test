"""
Report generation - Due Diligence Memo.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from jinja2 import Environment, FileSystemLoader

from .schemas import ReportInput, ReportOutput

# Template directory
TEMPLATE_DIR = Path(__file__).parent.parent / "templates"


def generate_report(input_data: ReportInput) -> ReportOutput:
    """
    Generate a due diligence memo from collected findings.

    Args:
        input_data: ReportInput with all findings

    Returns:
        ReportOutput with path to generated report
    """
    env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))
    template = env.get_template("memo_template.md")

    # Prepare template context
    context = {
        "company_name": input_data.company_name,
        "deal_overview": input_data.deal_overview,
        "assets": input_data.assets,
        "flood_risks": input_data.flood_risks,
        "transition_risk": input_data.transition_risk,
        "biodiversity": input_data.biodiversity,
        "esg_gaps": input_data.esg_gaps,
        "red_flags": input_data.red_flags,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }

    # Render markdown
    markdown_content = template.render(context)

    # Determine output path
    output_path = Path(input_data.output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if input_data.output_format == "markdown" or input_data.output_format == "html":
        # Save markdown
        md_path = output_path.with_suffix(".md")
        with open(md_path, "w") as f:
            f.write(markdown_content)

        if input_data.output_format == "html":
            # Convert to HTML
            html_path = output_path.with_suffix(".html")
            html_content = _markdown_to_html(markdown_content)
            with open(html_path, "w") as f:
                f.write(html_content)
            final_path = str(html_path)
        else:
            final_path = str(md_path)

    elif input_data.output_format == "docx":
        # Convert to DOCX
        docx_path = output_path.with_suffix(".docx")
        _markdown_to_docx(markdown_content, docx_path)
        final_path = str(docx_path)

    else:
        raise ValueError(f"Unsupported format: {input_data.output_format}")

    return ReportOutput(
        report_path=final_path,
        format=input_data.output_format,
        generated_at=datetime.now().isoformat(),
    )


def _markdown_to_html(markdown_content: str) -> str:
    """Convert markdown to standalone HTML."""
    import markdown

    html_body = markdown.markdown(markdown_content, extensions=["tables", "fenced_code"])

    html_template = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>CADI Due Diligence Memo</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{ color: #1a365d; border-bottom: 2px solid #2b6cb0; padding-bottom: 10px; }}
        h2 {{ color: #2c5282; margin-top: 30px; }}
        h3 {{ color: #2d3748; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #cbd5e0; padding: 10px; text-align: left; }}
        th {{ background-color: #edf2f7; }}
        .risk-high {{ color: #c53030; font-weight: bold; }}
        .risk-critical {{ color: #c53030; font-weight: bold; background: #fed7d7; padding: 2px 6px; }}
        .risk-medium {{ color: #d69e2e; }}
        .risk-low {{ color: #38a169; }}
        .red-flag {{ background: #fff5f5; border-left: 4px solid #c53030; padding: 10px; margin: 10px 0; }}
        .esg-gap {{ background: #fffaf0; border-left: 4px solid #d69e2e; padding: 10px; margin: 10px 0; }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""

    return html_template


def _markdown_to_docx(markdown_content: str, output_path: Path) -> None:
    """Convert markdown to DOCX using python-docx."""
    from docx import Document

    doc = Document()

    # Simple conversion - paragraphs and headers
    for line in markdown_content.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.strip():
            doc.add_paragraph(line)

    doc.save(str(output_path))
